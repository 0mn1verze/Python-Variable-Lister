""" 
Varlister

Created by Evan Fung, with help from Daniel Wright

Requirements: 
    python-docx
    docxs
"""


from typing import List
from collections import defaultdict
from pathlib import Path
import ast
import docx
import os
import codecs
import cProfile

"""========Config========"""


# Name and Path of output
OUTPUTNAME = 'Variable List'
OUTPUTPATH = os.getcwd()

# Names of files to ignore (Idea from Daniel)
IGNORE = []

# Add function names and return types using format -> [func-name]: [return-type]
USER_DEFINED_RETURN_TYPES = {'len': 'int',
                             'list': 'list',
                             'range': 'int',
                             'zip': 'list'}

"""=========Code========="""
# Heading for the documents
HEADING = ['Name', 'Type', 'Scope', 'Description', 'Line', 'Location']

# Containers
CONTAINERS = [ast.List, ast.Tuple, ast.Set]

# Types of variables with defined classes in ast
DEFINED_TYPES = {ast.List: "list", ast.Dict: "dict",
                 ast.Set: "set", ast.Tuple: "tuple",
                 ast.DictComp: "dict", ast.ListComp: "list",
                 ast.SetComp: "set", ast.GeneratorExp: "generator",
                 ast.JoinedStr: "str"}

# Which types to prioritise when guessing types
TYPE_PRIORITY = ['list', 'set', 'tuple',
                 'dict', 'generator', 'str', 'int', 'float']

# Defined return types for specific moderators
DEFINED_OPERATORS = {ast.Mod: "int",
                     ast.FloorDiv: "int", ast.Div: "float"}


class Node(object):
    def __init__(self, name: str = '*', func: str = None, cls: str = None, file: Path = None) -> None:
        """Node object

        Args:
            name (str, optional): name of the node. Defaults to '*'.
            func (str, optional): function name. Defaults to None.
            cls (str, optional): class name. Defaults to None.
            file (Path, optional): file path. Defaults to None.
        """
        self.name: str = name  # Scope name
        self.func: str = func  # Func name
        self.cls: str = cls  # Class name
        self.file: Path = file  # File name

        self.parent: object = None  # Parent node
        self.children: list = []  # Child nodes
        self.var_list: list[Variable] = []  # List of variables

    def __str__(self, syntax: str = "+- ", markers: List[str] = []) -> str:
        """_summary_

        Args:
            syntax (str, optional): syntax in front of each node. Defaults to "+- ".
            markers (list, optional): markers in front of parent node. Defaults to [].

        Returns:
            str: string representation of the tree starting from start node
        """
        # Init the two types of syntax in front of the node
        empty_str: str = " " * len(syntax)
        connect_str: str = "|" + empty_str[:-1]

        # Find the depth of the child relative to the starting node
        level: int = len(markers)

        # Return the two types of syntax with respect to marker[-1]
        def mapper(d: bool) -> str: return connect_str if d else empty_str
        marker: str = "".join(map(mapper, markers[:-1]))
        marker += syntax if level > 0 else ""

        # Representation of current node
        ret: str = f"{marker}{self.name}\n"

        # Prints the variables of the node
        for i, var in enumerate(self.var_list):
            if len(markers) > 0:
                var_marker = marker.replace(
                    syntax, '') + mapper(markers[-1]) + syntax
            else:
                var_marker = ''
            ret += f'{var_marker}{var.name}: {var.type}\n'

        # Recursive call the child string and concatenate the entire tree
        for i, child in enumerate(self.children):
            isLast: bool = i == len(self.children) - 1
            ret += child.__str__(syntax, [*markers, not isLast])

        return ret

    def add(self, name: str, cls: str = None, func: str = None, file: Path = None) -> object:
        """Adds a new child node to the current node

        Args:
            name (str): name of the node
            cls (str, optional): name of the class. Defaults to None.
            func (str, optional): name of the function. Defaults to None.
            file (Path, optional): path of the file. Defaults to None.

        Returns:
            Node: new node created
        """
        # Create new node, set its parent to the current node and append it to the current node
        new_node = Node(name, func, cls, file)
        self.children.append(new_node)
        new_node.parent = self
        return new_node

    def add_var(self, name: str, type: str, lineno: int):
        """Add Var to Node

        Args:
            name (str): name of the variable
            type (str): type of the variable
        """
        # Append the variable to the variable list
        self.var_list.append(Variable(name, self, lineno, type))

    def get_var(self, name: str, recurse: bool = False):
        """Get Var from ID

        Args:
            name (str): name of the variable

        Returns:
            Variable: variable with the same name
        """
        # For each variable in the variable list, check if its has the search name
        for var in self.var_list:
            if var.name == name:
                if var.type:
                    return var
        # For each children in the current node, check if it has the variable
        if recurse:
            for child in self.children:
                var = child.get_var(name, recurse)
                if var:
                    return var
        return None

    def get_return_type(self, name: str):
        """Get Return Type from Function Name

        Args:
            name (str): name of function

        Returns:
            str: return type of function
        """
        # For each children in the current node, get the return type
        for child in self.children:
            if child.func == name:
                return child.get_var("ReturnType")
        # For each children in the current node, get the return type of each of its children
        for child in self.children:
            return child.get_return_type(name)


class Variable(object):
    """Variable

    Stores the name, type and the scope data of a variable
    """
    varlist = defaultdict(list)

    def __init__(self, name: str, node: Node, lineno: int, type: str = None) -> None:
        self.name: str = name
        self.type: str = type
        self.lineno: int = lineno
        c: str = str(node.cls) if node.cls else ''
        f: str = str(node.func) if node.func else ''
        self.locations: List[str] = [f'{c} {f}']
        self.node: Node = node
        Variable.varlist[node.file.name].append(self)

    def __repr__(self):
        return self.name


class Visitor(ast.NodeVisitor):
    """Node Visitor

    Visits Nodes in the AST Object and find variable names and types
    """

    def __init__(self) -> None:
        self.global_node = Node()
        self.curr_node = self.global_node
        self.file_node = None
        self.curr_class = None

    def __str__(self) -> str:
        return str(self.global_node)

    def get_ID(self, node: ast.AST, include_subscript: bool = False) -> List[str] | str:
        """Get ID (Finds the name of the variable)

        Args:
            node (ast.AST): general ast node
            include_subscript (bool): whether to include subscript in the ID

        Returns:
            List[str]: full name of the variable
        """
        if any(isinstance(node, container) for container in CONTAINERS):
            # Returns the variable names inside the container
            return [self.get_ID(var) for var in node.elts]
        elif isinstance(node, ast.Name):
            # Returns the variable name
            return node.id
        elif isinstance(node, ast.Attribute):
            # Returns the attribute name
            return self.get_attr_ID(node)
        elif isinstance(node, ast.Call):
            # Returns the function name
            return self.get_call_ID(node)
        elif isinstance(node, ast.Subscript):
            # Returns the subscript name
            return self.get_subscript_ID(node, include_subscript)

    def get_attr_ID(self, node: ast.Attribute, include_subscript: bool = False) -> str:
        """Get Attribute ID (Recursively finds the name of the attribute)

        Args:
            node (ast.Attribute): node with attribute class
            include_subscript (bool): whether to include subscript in the ID

        Returns:
            str: full name of the attribute
        """
        val = node.value
        attr = node.attr
        if isinstance(val, ast.Name):
            # Concatenates the variable name and the attribute
            return f'{val.id}.{attr}'
        elif isinstance(val, ast.Attribute):
            # Concatenates the name of the attribute to the end of the name
            return f'{self.get_attr_ID(val)}.{attr}'
        elif isinstance(val, ast.Call):
            # Concatenates the name of the attribute to the end of the function name
            return f'{self.get_call_ID(val)}.{attr}'
        elif isinstance(val, ast.Subscript):
            # Concatenates the name of the attribute to the end of the subscript
            return f'{self.get_subscript_ID(val, include_subscript)}.{attr}'

    def get_call_ID(self, node: ast.Call) -> str:
        """Get Function ID (Finds the function name)

        Args:
            node (ast.Call): node with call class

        Returns:
            str: full name of the function
        """
        func = node.func
        # Gets the name of the function
        return f'{self.get_ID(func)}()'

    def get_subscript_ID(self, node: ast.Subscript, include_subscript: bool = False) -> str:
        """Get Subscript ID (Finds the subscript name)

        Args:
            node (ast.Subscript): node with subscript class
            include_subscript (bool): whether to include subscript in the ID

        Returns:
            str: full name of the subscript
        """
        val = node.value
        # Include the subscript if the argument is set to True
        if include_subscript:
            names = self.get_ID(node.slice)
            subscript = ', '.join(list(map(str, names))) if isinstance(
                names, list) else names
            return f'{self.get_ID(val, include_subscript)}[{subscript}]'
        return f'{self.get_ID(val)}[]'

    def get_type(self, node: ast.AST) -> str:
        """Get Variable Types

        Args:
            node (ast.AST): value of assignment

        Returns:
            str: type of variable
        """
        if not node:
            return None
        # Loops through the defined types and checking if the node is one of the types
        for ast_type, var_type in DEFINED_TYPES.items():
            if isinstance(node, ast_type):
                return var_type
        else:
            # If node is a constant then return the type of the constant
            if isinstance(node, ast.Constant):
                return type(node.value).__name__
            # If node is a name then return the variable name
            if isinstance(node, ast.Name):
                name = self.get_ID(node)
                var = self.curr_node.get_var(name)
                if not var and name != 'self':
                    var_type = name
                return var_type
            # If node is an attribute then return the attribute name
            if isinstance(node, ast.Attribute):
                return self.get_attr_ID(node)
            # If node is a function then return the function name
            if isinstance(node, ast.Call):
                name = self.get_call_ID(node)
                if 'self.' in name:
                    return self.curr_class.get_return_type(name.replace('self.', ''))
                # If the return type is defined by the user then use the return type
                code = name.replace('()', '')
                if code in USER_DEFINED_RETURN_TYPES:
                    return USER_DEFINED_RETURN_TYPES[code]
                return name
            # If node is a subscript then return the subscript name
            if isinstance(node, ast.Subscript):
                return self.get_subscript_ID(node)
            # If node is a binop then return the type of one of the variables with priorities
            if isinstance(node, ast.BinOp):
                # If the operation in the node can only return a specific type then return the variable type
                for ast_type, var_type in DEFINED_OPERATORS.items():
                    if isinstance(node.op, ast_type):
                        return var_type

                # Get the list of types for each variable in the binary operation
                types = self.get_binop_types(node)

                # Initialise the variables for finding the correct type based on a type priority
                var_type = types[0]
                var_type_idx = 9999

                # Find the correct type based on a type priority
                for t in types:
                    if t in TYPE_PRIORITY:
                        idx = TYPE_PRIORITY.index(t)
                        if idx < var_type_idx:
                            var_type_idx = idx
                            var_type = t
                return var_type
            # If node is a inline if expression then find the type of the test variable
            if isinstance(node, ast.IfExp):
                return self.get_type(node.test)
            # If node is a boolean operation then return boolean as the type
            if isinstance(node, ast.BoolOp):
                return "bool"
            # If node is a unary operation then return the type of the operand
            if isinstance(node, ast.UnaryOp):
                return self.get_type(node.operand)
            # If node is a comparison then return boolean as the type
            if isinstance(node, ast.Compare):
                return "bool"
            if isinstance(node, ast.Await):
                return self.get_type(node.value)
            if isinstance(node, ast.Lambda):
                return "lambda function"
        print(ast.dump(node, indent=2))

    def get_binop_types(self, node: ast.BinOp) -> List[str]:
        """Get Binary Operation Type

        Args:
            node (ast.BinOP): node with binop class

        Returns:
            List[str]: list of types in the binary operation
        """
        # returns the type of the left and right of the binary operation
        return [self.get_type(node.left), self.get_type(node.right)]

    def get_var(self, name: str) -> Variable:
        """Get Variable From Node

        Args:
            name (str): name of variable

        Returns:
            Variable: variable with the corresponding name
        """
        # If self. is in the name then find the variable from the current class
        if name:
            if 'self.' in name:
                var = self.curr_class.get_var(name, True)
                if var:
                    c = str(self.curr_node.cls) if self.curr_node.cls else ''
                    f = str(self.curr_node.func) if self.curr_node.func else ''
                    var.locations.append(f'{c} {f}')
                    return var
                return None
            else:
                # If the variable exists then return the variable or else search the global node
                var = self.curr_node.get_var(name)
                if var:
                    return var
                # Try to find the variable in the global scope only
                return self.file_node.get_var(name)

    def handle_comp(self, node: ast.ListComp | ast.SetComp | ast.DictComp | ast.GeneratorExp):
        """Handle Comprehensions and generators (Idea from Daniel)

        Args:
            node (ast.ListComp | ast.SetComp | ast.DictComp | ast.GeneratorExp): generator and comprehension nodes
        """
        # For every comprehension in the node, find the variable and its type
        for comp in node.generators:
            names = self.get_ID(comp.target)
            names = [names] if isinstance(names, str) else names
            # For every name in the names from the target, get the variable type
            for name in names:
                var_type = self.get_type(comp.iter)
                var = self.get_var(var_type)
                if var:
                    var_type = var.type
                self.curr_node.add_var(name, var_type, node.lineno)

    def handle_args(self, node: ast.FunctionDef | ast.Lambda):
        """Handle Arguments from Functions 

        Args:
            node (ast.FunctionDef | ast.Lambda): function definition or lambda nodes
        """
        # Initialising variables
        args = node.args.args
        defaults = node.args.defaults
        # Setting defaults to align with the corresponding variables
        defaults = [None] * (len(args) - len(defaults)) + defaults
        kwargs = node.args.kwonlyargs
        kwdefaults = node.args.kw_defaults

        # For each argument find its type
        for idx, arg in enumerate(args):
            name = arg.arg
            if name != 'self':
                if 'annotation' in arg._fields:
                    arg_type = self.get_ann_type(arg.annotation)
                else:
                    arg_type = self.get_type(defaults[idx])
                self.curr_node.add_var(name, arg_type, node.lineno)

        # For each keyword argument find its type
        for idx, kwarg in enumerate(kwargs):
            name = kwarg.arg
            if 'annotation' in kwarg._fields:
                arg_type = self.get_ann_type(kwarg.annotation)
            else:
                arg_type = self.get_type(kwdefaults[idx])
            self.curr_node.add_var(name, arg_type, node.lineno)

    def get_ann_type(self, ann: ast.Subscript | ast.Attribute | ast.Name) -> str:
        """Handle Annotations

        Args:
            ann (ast.Subscript | ast.Attribute | ast.Name): annotation

        Returns:
            str: annotation type
        """
        # If annotation is a name then return the name
        if isinstance(ann, ast.Name):
            return ann.id
        # If annotation is a subscript then return the subscript
        elif isinstance(ann, ast.Subscript):
            return self.get_subscript_ID(ann, include_subscript=True)
        # If annotation is an attribute then return the name of the attribute
        elif isinstance(ann, ast.Attribute):
            return self.get_attr_ID(ann, include_subscript=True)

    def visit_ClassDef(self, node: ast.ClassDef) -> None:
        """Visit Class Definitions

        Args:
            node (ast.ClassDef): class definition node
        """
        # Add a new node to the current node and set the current node to the new node
        self.curr_node = self.curr_node.add(
            node.name, node.name, None, self.curr_node.file)
        self.curr_class = self.curr_node
        super().generic_visit(node)
        # Restore the previous node
        self.curr_node = self.curr_node.parent

    def visit_FunctionDef(self, node: ast.FunctionDef) -> None:
        """Visit Function Definitions

        Args:
            node (ast.FunctionDef): function definiton node
        """
        # Add a new node to the current node and set the current node to the new node
        self.curr_node = self.curr_node.add(
            node.name, self.curr_node.cls, node.name, self.curr_node.file)
        # Handle function arguments
        self.handle_args(node)
        # Add return type if there is annotation
        if node.returns:
            ann_type = self.get_ann_type(node.returns)
            if ann_type:
                self.curr_node.add_var("ReturnType", ann_type, node.lineno)
        super().generic_visit(node)
        # Restore the previous node
        self.curr_node = self.curr_node.parent

    def visit_Return(self, node: ast.Return) -> None:
        """Visit Return Nodes

        Args:
            node (ast.Return): return nodes
        """
        super().generic_visit(node)
        # After visiting the return node add the return type to the scope
        name = "ReturnType"
        # If the returntype is already found then ignore the return node
        var = self.get_var(name)
        var_type = self.get_type(node.value)
        # If the type is none then do not create a new returntype variable
        if not var and var_type and name != 'self':
            self.curr_node.add_var(name, var_type, node.lineno)

    def visit_Yield(self, node: ast.Yield) -> None:
        """Visit Yield Nodes

        Args:
            node (ast.Yield): Yield nodes
        """
        super().generic_visit(node)
        # After visiting the yield node add the yield type to the scope
        name = "YieldType"
        # If the yieldtype is already found then ignore the yield node
        var = self.get_var(name)
        var_type = self.get_type(node.value)
        # If the type is none then do not create a new yieldtype variable
        if not var and var_type and name != 'self':
            self.curr_node.add_var(name, var_type, node.lineno)

    def visit_YieldFrom(self, node: ast.YieldFrom) -> None:
        """Visit YieldFrom Nodes

        Args:
            node (ast.YieldFrom): YieldFrom nodes
        """
        super().generic_visit(node)
        # After visiting the yield node add the yield type to the scope
        name = "YieldType"
        # If the yieldtype is already found then ignore the yield node
        var = self.get_var(name)
        var_type = self.get_type(node.value)
        # If the type is none then do not create a new yieldtype variable
        if not var and var_type and name != 'self':
            self.curr_node.add_var(name, var_type, node.lineno)

    def visit_Assign(self, node: ast.Assign) -> None:
        """Visit Assign Nodes

        Args:
            node (ast.Assign): assign nodes
        """
        # Names can be a string or a list of strings
        names = self.get_ID(node.targets[0])

        # Converts names into a list of strings
        names = [names] if isinstance(names, str) else names

        # Loops through the names and finds the type of the element assigned to it
        for idx, name in enumerate(names):
            # Get variable from the current node if it exists
            var = self.get_var(name)
            # If the variable exists ignore it
            if not var and name != 'self':
                # If the value is a container then find the type of the element at the same index as the name
                if len(node.targets) > 1 and any(isinstance(node.value, t) for t in CONTAINERS) and len(node.value.elts) > 0:
                    var_type = self.get_type(node.value.elts[idx])
                # If the value is not a container then get the type of the node
                else:
                    var_type = self.get_type(node.value)
                self.curr_node.add_var(name, var_type, node.lineno)
        super().generic_visit(node)

    def visit_AugAssign(self, node: ast.AugAssign) -> None:
        """Visit Augment Assign Nodes

        Args:
            node (ast.AugAssign): augment assign nodes
        """
        # Get the name and type of the variable and add a new variable if it doesn't exist already
        name = self.get_ID(node.target)
        var = self.get_var(name)
        if not var and name != 'self':
            var_type = self.get_type(node.value)
            self.curr_node.add_var(name, var_type, node.lineno)
        super().generic_visit(node)

    def visit_AnnAssign(self, node: ast.AnnAssign) -> None:
        """Visit Annotated Assign Nodes

        Args:
            node (ast.AnnAssign): annotated assign nodes
        """
        # Get the name and annotation of the node
        name = self.get_ID(node.target)
        ann = node.annotation

        # If var does not exist in the node tree then determine the name of the annotation
        var = self.get_var(name)
        if not var and name != 'self':
            var_type = self.get_ann_type(ann)
            self.curr_node.add_var(name, var_type, node.lineno)
        super().generic_visit(node)

    def visit_ListComp(self, node: ast.ListComp) -> None:
        self.handle_comp(node)
        super().generic_visit(node)

    def visit_SetComp(self, node: ast.SetComp) -> None:
        self.handle_comp(node)
        super().generic_visit(node)

    def visit_DictComp(self, node: ast.DictComp) -> None:
        self.handle_comp(node)
        super().generic_visit(node)

    def visit_GeneratorExp(self, node: ast.GeneratorExp) -> None:
        self.handle_comp(node)
        super().generic_visit(node)

    def visit_For(self, node: ast.For) -> None:
        """Visit For loops

        Args:
            node (ast.For): for loop nodes
        """
        # Get the names and types of the variables and add new variables if it doesn't exist already
        names = self.get_ID(node.target)
        names = [names] if isinstance(names, str) else names
        for name in names:
            var_type = self.get_type(node.iter)
            self.curr_node.add_var(name, var_type, node.lineno)
        super().generic_visit(node)

    def visit_With(self, node: ast.With) -> None:
        """Visit With Statements

        Args:
            node (ast.With): with statement nodes
        """
        # Get the names and types of the variables and add new variables if it doesn't exist already
        withitem = node.items[0]
        name = self.get_ID(withitem.optional_vars)
        var_type = self.get_type(withitem.context_expr)
        self.curr_node.add_var(name, var_type, node.lineno)
        super().generic_visit(node)

    def visit_Lambda(self, node: ast.Lambda) -> None:
        self.handle_args(node)
        super().generic_visit(node)

    def parse_file(self, file: Path):
        """Parse File

        Args:
            file (Path): path of file
        """
        # Adds the file as a new node to the parent node
        self.file_node = self.global_node.add(file, file=file)
        self.curr_node = self.file_node

        # Opens the file, compiles it and reads the variables
        with codecs.open(file, 'r', 'utf-8') as f:
            print(f"Parsing File: {file.name}...")
            compilation = ast.parse(f.read())
            self.visit(compilation)

        self.file_node = None

    def parse_cwd(self):
        """Parses all the py files in the current working directory and all its subdirectories
        """
        files = list(Path.cwd().rglob('*.py'))
        for f in files:
            if f == Path(__file__) or f.name in IGNORE:
                continue
            self.parse_file(f)

    def create_document(self):
        """Create Variable Listings Document 
        """
        # Creates document
        print("Creating Document...")
        document = docx.Document()
        document.add_heading('Variable List', level=1)

        # Initialises varlist
        varlist = Variable.varlist
        # Loops through the files
        for file_name, vars in varlist.items():
            # Creates a new table for each file
            print(
                f'Creating Table for {file_name}... ({len(vars)} variables including return types)')
            document.add_heading(str(file_name), level=2)
            rows = len(vars)+1
            cols = len(HEADING)
            table = document.add_table(rows=rows, cols=cols)
            table.style = "TableGrid"
            cells = table._cells

            # Fill in the header
            for col, col_heading in enumerate(HEADING):
                cells[col].text = col_heading

            # Fill in the table with variables and corresponding data
            for row, var in enumerate(vars, 1):
                if var.name and var.name != "ReturnType" and var.name != "YieldType":
                    node = var.node
                    idx = row * cols
                    cells[idx].text = str(var.name)
                    cells[1 +
                          idx].text = str(var.type) if var.type else 'Unknown'
                    cells[2 + idx].text = 'Global' if not node.func and not node.cls else 'Local'
                    cells[4 + idx].text = str(var.lineno)
                    cells[5 + idx].text = ', '.join(var.locations)
        # Save Document
        print('Saving Document...')
        document.save(f"{OUTPUTPATH}/{OUTPUTNAME}.docx")
        path = OUTPUTPATH if OUTPUTPATH != os.getcwd() else 'current directory'
        print(f'Document saved to {path}!')


if __name__ == "__main__":
    visitor = Visitor()
    visitor.parse_cwd()
    visitor.create_document()
