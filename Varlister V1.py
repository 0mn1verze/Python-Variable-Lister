import ast
from collections import defaultdict
import csv
import glob
import os
from pathlib import Path
from itertools import groupby


class Scope(object):
    vars = defaultdict(str)

    def __init__(self, name, function=False):
        self.name = name
        self.children = []
        self.vars = defaultdict(str)
        self.function = function
        self.classname = None
        self.parent = None

    def __repr__(self):
        return self.name

    def add(self, name, function=False, classname=None):
        new_scope = Scope(name, function)
        self.children.append(new_scope)
        new_scope.parent = self
        if classname:
            new_scope.classname = classname
        else:
            new_scope.classname = self.classname
        return new_scope

    def print(self, indent=0):
        padding = " " * indent
        print(padding + self.name + str(self.classname if self.classname else ''))
        indent += 4
        padding = " " * indent
        for var, item in self.vars.items():
            if item:
                if isinstance(item[0], type):
                    print(
                        padding + f'{var} (line {item[1]}): {item[0].__name__}')
                else:
                    print(
                        padding + f'{var} (line {item[1]}): {str(item[0])}')
            else:
                print(
                    padding + f'{var}: None')
        for child in self.children:
            child.print(indent)

    def get_full_var_list(self, location=''):
        scope = 'local' if self.name != '*' and '.py' not in self.name else 'global'
        if self.name != '*':
            location = f'{location} {self.name}'
        for var, item in self.vars.items():
            if var == 'ReturnType':
                continue
            varname = var + ' ' + str(self.classname if self.classname else '')
            if 'self.' in var and varname in Scope.vars:
                Scope.vars[varname][4] += ', ' + location
                Scope.vars[varname][3] += ', ' + f'{item[1]}'
            elif item:
                if item[0] and item[1]:
                    if isinstance(item[0], type):
                        Scope.vars[varname] = [item[0].__name__,
                                               scope, '', f'{item[1]}',  location]
                    else:
                        Scope.vars[varname] = [
                            item[0], scope, '', f'{item[1]}',  location]
                elif item[1]:
                    Scope.vars[varname] = ['unknown',
                                           scope, '', f'{item[1]}',  location]
                else:
                    Scope.vars[varname] = ['unknown',
                                           scope, '', f'unknown',  location]
        for child in self.children:
            child.get_full_var_list(location)

    def add_var(self, name, lineno, vartype=None):
        if name not in self.vars:
            if vartype:
                self.vars[name] = [vartype, lineno]
            elif lineno:
                self.vars[name] = [None, lineno]
            else:
                self.vars[name] = [None, None]

    def update_var(self, name, lineno=None, vartype=None):
        if name in self.vars:
            if self.vars[name]:
                self.vars[name][0] = vartype
            else:
                self.vars[name] = [vartype, lineno]
        else:
            return False

    def update_var_with_varname(self, varname, vartype):
        for var, vartype in self.vars.items():
            if vartype == varname:
                self.vars[var][0] = vartype
        for child in self.children:
            child.update_var_with_varname(varname, vartype)

    def update_return_type(self, funcname, returntype):
        for var, vartype in self.vars.items():
            if vartype:
                if vartype[0] == funcname:
                    self.vars[var][0] = returntype
                    self.update_var_with_varname(var, returntype)
        for child in self.children:
            child.update_return_type(funcname, returntype)

    def find_return_type(self, funcname):
        for child in self.children:
            if child.name == funcname:
                if 'ReturnType' in child.vars:
                    return child.vars['ReturnType'][0]
                else:
                    return child.find_return_type(funcname)
        else:
            return False

    def find_class_var_type(self, name):
        for child in self.parent.children:
            if name in child.vars:
                return child.vars[name][0]
        else:
            return False


class VarLister(ast.NodeVisitor):
    def __init__(self, scope, global_scope):
        self.scope = scope
        self.global_scope = global_scope

    def visit_ClassDef(self, node):
        self.scope = self.scope.add(node.name, False, node.name)
        super().generic_visit(node)
        self.scope = self.scope.parent

    def get_subscript_type(self, node):
        var = node.value
        if isinstance(var, ast.Name):
            return var.id
        elif isinstance(var, ast.Subscript):
            if isinstance(slice, ast.Name):
                return self.get_subscript_type(var) + f'[{var.slice.id}]'
            else:
                return self.get_subscript_type(var)
        else:
            return False

    def visit_FunctionDef(self, node):
        self.scope = self.scope.add(node.name, True)
        for idx, arg in enumerate(node.args.args):
            if arg.arg != 'self':
                if isinstance(arg.annotation, ast.Name):
                    self.scope.add_var(arg.arg, arg.lineno, arg.annotation.id)
                if isinstance(arg.annotation, ast.Subscript):
                    self.scope.add_var(arg.arg, arg.lineno,
                                       self.get_subscript_type(arg.annotation))
                default_idx = len(node.args.args) - len(node.args.defaults)
                if idx >= default_idx:
                    self.scope.add_var(
                        arg.arg, arg.lineno, self.get_var_type(node.args.defaults[idx - default_idx]))
        super().generic_visit(node)
        if isinstance(node.returns, ast.Name):
            self.scope.add_var('ReturnType', node.lineno, node.returns.id)
        elif isinstance(node.returns, ast.Constant):
            self.scope.add_var('ReturnType', node.lineno,
                               type(node.returns.value))
        elif isinstance(node.returns, ast.Subscript):
            self.scope.add_var('ReturnType', node.lineno,
                               self.get_subscript_type(node.returns))
        elif isinstance(node.body[-1], ast.Return):
            return_value = node.body[-1].value
            vartype = self.get_var_type(return_value)
            self.scope.add_var('ReturnType', node.body[-1].lineno, vartype)
            self.global_scope.update_return_type(node.name + '()', vartype)
            self.global_scope.update_return_type(
                'self.' + node.name + '()', vartype)
        self.scope = self.scope.parent

    def get_binop_names(self, var):
        names = []
        for val in [var.left, var.right]:
            if isinstance(val, ast.BinOp):
                names.extend(self.get_binop_names(val))
            elif isinstance(val, ast.Attribute):
                names.append(self.get_attr_name(val))
            elif isinstance(val, ast.Name):
                names.append(val.id)
            elif isinstance(var, ast.Call):
                function = var.func
                if isinstance(function, ast.Attribute):
                    names.append(self.get_attr_name(function) + '()')
                elif isinstance(function, ast.Name):
                    names.append(function.id + '()')
        return names

    def get_binop_types(self, var):
        types = []
        for val in [var.left, var.right]:
            if isinstance(val, ast.BinOp):
                types.extend(self.get_binop_types(val))
            else:
                types.append(self.get_var_type(val))
        return types

    def get_attr_name(self, var, is_subscript=False):
        name = ''
        if is_subscript:
            ending = '[]'
        else:
            ending = '.' + var.attr
        if isinstance(var.value, ast.Attribute):
            name += self.get_attr_name(var.value) + ending
        elif isinstance(var.value, ast.Name):
            name = var.value.id + ending
        elif isinstance(var.value, ast.Call):
            name = self.get_func_name(var.value) + ending
        elif isinstance(var.value, ast.Subscript):
            name = self.get_attr_name(var.value, True) + ending
        return name

    def get_func_name(self, var):
        function = var.func
        if isinstance(function, ast.Attribute):
            vartype = self.get_attr_name(function) + '()'
            if isinstance(function.value, ast.Constant):
                vartype = type(function.value.value)
        elif isinstance(function, ast.Name):
            returntype = self.global_scope.find_return_type(function.id)
            if returntype:
                vartype = returntype
            else:
                if function.id in ['dict', 'set', 'tuple', 'list', 'float', 'int', 'str', 'bool']:
                    vartype = function.id
                else:
                    vartype = function.id + '()'
        else:
            vartype = None
        return vartype

    def get_prefered_type(self, types):
        type_priority = [tuple, list, str, int]
        for t in type_priority:
            if t in types:
                vartype = t
                break
        else:
            vartype = types[0]
        return vartype

    def get_var_type(self, var):
        vartype = None

        if isinstance(var, ast.Constant):
            vartype = type(var.value)
        elif isinstance(var, ast.List):
            vartype = list
        elif isinstance(var, ast.Tuple):
            vartype = tuple
        elif isinstance(var, ast.Dict):
            vartype = dict
        elif isinstance(var, ast.JoinedStr):
            vartype = str
        elif isinstance(var, ast.Name):
            if var.id in self.scope.vars:
                if self.scope.vars[var.id]:
                    vartype = self.scope.vars[var.id][0]
                else:
                    vartype = None
            else:
                vartype = None
        elif isinstance(var, ast.BinOp):
            types = self.get_binop_types(var)
            vartype = self.get_prefered_type(types)
        elif isinstance(var, ast.Attribute):
            vartype = self.get_attr_name(var)
        elif isinstance(var, ast.Call):
            vartype = self.get_func_name(var)
        elif isinstance(var, ast.IfExp):
            types = [self.get_var_type(
                var.body), self.get_var_type(var.orelse)]
            vartype = self.get_prefered_type(types)
        elif isinstance(var, ast.UnaryOp):
            vartype = self.get_var_type(var.operand)
        elif isinstance(var, ast.Subscript):
            if isinstance(var.value, ast.Call):
                if isinstance(var.slice, ast.Slice):
                    vartype = self.get_func_name(var.value)
                else:
                    vartype = self.get_func_name(var.value) + '[]'
            elif isinstance(var.value, ast.Attribute):
                vartype = self.get_attr_name(var.value) + '[]'
        return vartype

    def update_subscript_type(self, node):
        var = node.value
        slice = node.slice
        if isinstance(var, ast.Name):
            if var.id in self.scope.vars:
                if self.get_var_type(slice) == int:
                    self.scope.update_var(var.id, node.lineno, list)
            if isinstance(slice, ast.Name):
                vartype = None
                if var.id in self.scope.vars:
                    vartype = self.scope.vars[var.id]
                if vartype == list:
                    self.scope.update_var(slice.id, node.lineno, int)
                    return True
            elif isinstance(slice, ast.BinOp):
                names = self.get_binop_names(slice)
                for name in names:
                    if self.scope.vars[name] == list:
                        self.scope.update_var(name, node.lineno, int)
                        return True
            return False
        elif isinstance(var, ast.Subscript):
            if self.update_subscript_type(var):
                if isinstance(slice, ast.Name):
                    self.scope.update_var(slice.id, node.lineno, int)

    def visit_Subscript(self, node):
        self.update_subscript_type(node)
        super().generic_visit(node)

    def analyze_var(self, val, target):
        tuple_or_list = (ast.Tuple, ast.List)
        if isinstance(target, ast.Name):
            name = target.id
            vartype = self.get_var_type(val)
            self.scope.add_var(name, target.lineno, vartype)
        elif isinstance(target, tuple_or_list) and isinstance(val, tuple_or_list):
            for idx, var in enumerate(target.elts):
                vartype = None
                name = None
                if isinstance(var, ast.Name):
                    name = var.id
                    v = val.elts[idx]
                    vartype = self.get_var_type(v)
                    self.scope.add_var(name, target.lineno, vartype)
                elif isinstance(var, ast.Subscript):
                    if isinstance(var.value, ast.Name):
                        name = var.value.id
                        self.scope.update_var(name, target.lineno, list)
        elif isinstance(target, ast.Attribute):
            name = self.get_attr_name(target)
            if 'self.' in name:
                vartype = self.scope.find_class_var_type(name)
                if not vartype:
                    vartype = self.get_var_type(val)
            else:
                vartype = self.get_var_type(val)
            self.scope.add_var(name, target.lineno, vartype)

    def visit_Assign(self, node):
        target = node.targets[0]
        val = node.value
        self.analyze_var(val, target)
        super().generic_visit(node)

    def visit_AnnAssign(self, node):
        target = node.target
        ann = node.annotation
        vartype = None
        if isinstance(ann, ast.Name):
            vartype = ann.id
        elif isinstance(ann, ast.Subscript):
            vartype = self.get_subscript_type(ann)
        if isinstance(target, ast.Name):
            self.scope.add_var(target.id, target.lineno, vartype)
        elif isinstance(target, ast.Attribute):
            name = self.get_attr_name(target)
            self.scope.add_var(name, target.lineno, vartype)
        super().generic_visit(node)

    def visit_AugAssign(self, node):
        target = node.target
        val = node.value
        self.analyze_var(val, target)
        super().generic_visit(node)

    def visit_Call(self, node):
        super().generic_visit(node)


def main(use_csv=False):

    # input file names of the files
    files = list(Path.cwd().rglob('*.py'))

    print(f"Files read: {', '.join(map(lambda x: x.name, files))}")

    scope = Scope('*')
    for f in files:
        if f != os.path.relpath(__file__):
            with open(f, 'r') as file:
                parsed = ast.parse(file.read())
                current_scope = scope.add(str(f.name), True)
                varlister = VarLister(current_scope, scope)
                varlister.visit(parsed)
    fields = ['variable name', 'type', 'scope',
              'description', 'line', 'function']

    filename = 'variablelist.csv'

    scope.get_full_var_list()
    vars = []
    Scope.vars = {k: v for k, v in sorted(
        Scope.vars.items(), key=lambda x: x[1][4])}
    for var, desc in Scope.vars.items():
        vars.append([var.split(' ')[0], *desc])

    if use_csv:
        with open(filename, 'w') as csvfile:
            csvwriter = csv.writer(csvfile, lineterminator='\n')
            csvwriter.writerow(fields)

            csvwriter.writerows(vars)
    else:
        from docx import Document
        document = Document()
        document.add_heading(f'Variable list', level=1)
        for file, group in groupby(vars, lambda x: x[5].split(' ')[1]):
            varlist = list(group)
            document.add_heading(file, level=2)
            table = document.add_table(rows=len(varlist)+1, cols=6)
            table.style = "TableGrid"
            for col, fieldname in enumerate(fields):
                table.cell(0, col).text = fieldname
            for row, record in enumerate(varlist, 1):
                for col, fielddata in enumerate(record):
                    if col == 5:
                        fielddata = ' '.join(fielddata.split(' ')[2:])
                    table.cell(row, col).text = fielddata
        document.save(f"./Variable list.docx")


main(True)

a = [i for i in range(10)]
